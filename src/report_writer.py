"""
Word报告导出模块
生成专业格式的公路路况分析报告
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd


def set_cell_bg(cell, hex_color: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_borders(table):
    """为表格设置细边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_heading(doc, text: str, level: int):
    """添加标题"""
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(14)
    para.paragraph_format.line_spacing = 1.5


def add_body_text(doc, text: str):
    """添加正文段落"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    return para


def add_image(doc, img_path: str, caption: str = None, width_inches: float = 5.5):
    """添加图片和图注"""
    if not img_path or not os.path.exists(img_path):
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    try:
        run.add_picture(img_path, width=Inches(width_inches))
    except Exception as e:
        para.add_run(f'[图片加载失败: {e}]')
    
    if caption:
        cap_para = doc.add_paragraph(caption)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.runs[0]
        cap_run.font.size = Pt(14)
        cap_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        cap_run.font.name = '仿宋'
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')


def add_summary_table(doc, df: pd.DataFrame, title: str = None):
    """添加汇总统计表格"""
    if df.empty:
        return
    
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    cols = list(df.columns)
    table = doc.add_table(rows=len(df) + 1, cols=len(cols) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # 表头
    header_row = table.rows[0]
    set_cell_bg(header_row.cells[0], 'D6E4F0')
    header_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for j, col in enumerate(cols):
        cell = header_row.cells[j + 1]
        set_cell_bg(cell, 'D6E4F0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.size = Pt(12)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 数据行
    for i, (idx, row) in enumerate(df.iterrows()):
        tr = table.rows[i + 1]
        if i % 2 == 0:
            set_cell_bg(tr.cells[0], 'F5F5F5')
        # 行索引
        idx_cell = tr.cells[0]
        idx_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = idx_cell.paragraphs[0].add_run(str(idx))
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        
        for j, val in enumerate(row):
            cell = tr.cells[j + 1]
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(round(val, 2) if isinstance(val, float) else val))
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_paragraph()


def add_data_table(doc, data: list, headers: list, title: str = None):
    """添加数据表格"""
    if not data:
        return
    
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # 表头
    header_row = table.rows[0]
    for j, col in enumerate(headers):
        cell = header_row.cells[j]
        set_cell_bg(cell, 'D6E4F0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.size = Pt(12)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 数据行
    for i, row_data in enumerate(data):
        tr = table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = tr.cells[j]
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(round(val, 2) if isinstance(val, float) else val))
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_paragraph()


def generate_word_report(
    df: pd.DataFrame,
    output_path: str,
    title: str = None,
    summary: str = None,
    api_key: str = None,
    base_url: str = None,
    model: str = None,
    table_data: dict = None,
    chart_paths: dict = None,
    progress_callback=None
):
    """
    生成完整的Word格式分析报告
    
    Args:
        df: 数据DataFrame
        output_path: Word文件保存路径
        title: 报告标题
        summary: 报告摘要
        api_key: AI API密钥
        base_url: AI API基础URL
        model: AI模型名称
        table_data: 表格数据字典，包含各种分析结果
        chart_paths: 图表路径字典
        progress_callback: 进度回调函数
    """
    def _progress(msg):
        if progress_callback:
            progress_callback(msg)
    
    doc = Document()
    
    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    
    # ── 默认字体 ──
    doc.styles['Normal'].font.name = '仿宋'
    doc.styles['Normal'].font.size = Pt(14)
    doc.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 从df中提取信息
    if df.empty:
        return "数据为空，无法生成报告"
    
    # 获取县份信息
    counties = df['县份'].unique().tolist() if '县份' in df.columns else ['未知']
    county = counties[0] if len(counties) == 1 else '全部'
    
    # 获取年份信息
    years = sorted(df['年份'].unique().tolist()) if '年份' in df.columns else []
    latest_year = max(years) if years else datetime.now().year
    
    # 生成报告标题
    report_title = title or f'{county}公路路况分析报告（{min(years) if years else ""}—{max(years) if years else ""}年）'
    
    # 生成基本统计数据
    stats = {
        'total_segments': len(df),
        'total_length_km': df['路段长度km'].sum() if '路段长度km' in df.columns else 0,
        'PQI_mean': round(df['PQI'].mean(), 2) if 'PQI' in df.columns else 'N/A',
        'PCI_mean': round(df['PCI'].mean(), 2) if 'PCI' in df.columns else 'N/A',
        'RQI_mean': round(df['RQI'].mean(), 2) if 'RQI' in df.columns else 'N/A'
    }
    
    # 初始化AI客户端
    llm = None
    if api_key:
        try:
            from src.llm_writer import LLMClient
            llm = LLMClient(api_key=api_key, base_url=base_url, model=model)
        except Exception as e:
            print(f"AI初始化失败: {e}")
    
    # 生成报告预览文本
    preview_text = f"报告标题：{report_title}\n"
    preview_text += f"分析县份：{county}\n"
    preview_text += f"数据年份：{', '.join(str(y) for y in years) if years else '未知'}\n"
    preview_text += f"检测路段数：{stats['total_segments']}\n"
    preview_text += f"总里程：{stats['total_length_km']:.1f} km\n"
    preview_text += f"PQI均值：{stats['PQI_mean']}\n"
    preview_text += f"PCI均值：{stats['PCI_mean']}\n"
    preview_text += f"RQI均值：{stats['RQI_mean']}\n"
    
    # ═══════════════════════
    # 封面
    # ═══════════════════════
    _progress("正在生成封面...")
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(report_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_run.font.name = '仿宋'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_paragraph()
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f'基于{min(years) if years else ""}—{max(years) if years else ""}年路面技术状况检测数据')
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    sub_run.font.name = '仿宋'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime('%Y年%m月'))
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    date_run.font.name = '仿宋'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_page_break()
    
    # ═══════════════════════
    # 一、概述
    # ═══════════════════════
    _progress("正在生成概述...")
    add_heading(doc, '一、概述', 1)
    
    # 生成AI概述文字
    overview_text = None
    if llm and table_data:
        try:
            from src.llm_writer import generate_overview_text
            overview_text = generate_overview_text(stats, county, years, llm)
        except Exception as e:
            print(f"生成概述失败: {e}")
    
    if summary:
        add_body_text(doc, summary)
        if overview_text:
            add_body_text(doc, overview_text)
    elif overview_text:
        add_body_text(doc, overview_text)
    else:
        add_body_text(doc, f'本报告以{county}公路{min(years) if years else ""}—{max(years) if years else ""}年路面技术状况检测数据为基础，对路面技术状况指数（PQI、PCI、RQI）进行系统分析，评估路面技术状况等级分布，并对养护需求进行预测。')
    
    # 基本情况表
    add_heading(doc, '1.1 数据基本情况', 2)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_table)
    
    info_data = [
        ('分析县份', county),
        ('数据年份', ', '.join(str(y) for y in years) if years else '-'),
        ('检测路段数', f"{stats.get('total_segments', '-')} 段"),
        ('总里程', f"{stats.get('total_length_km', 0):.1f} km"),
        ('检测指标', 'PQI（路面质量指数）、PCI（路面状况指数）、RQI（行驶质量指数）'),
    ]
    for i, (key, val) in enumerate(info_data):
        row = info_table.rows[i]
        set_cell_bg(row.cells[0], 'D6E4F0')
        p0 = row.cells[0].paragraphs[0]
        run0 = p0.add_run(key)
        run0.bold = True
        run0.font.size = Pt(12)
        run0.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run0.font.name = '仿宋'
        run0._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run1 = row.cells[1].paragraphs[0].add_run(val)
        run1.font.size = Pt(12)
        run1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run1.font.name = '仿宋'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.add_paragraph()
    
    # ═══════════════════════
    # 二、路面技术状况年度趋势
    # ═══════════════════════
    _progress("正在生成趋势分析...")
    add_heading(doc, '二、路面技术状况年度趋势分析', 1)
    
    # 添加趋势图表
    if chart_paths and 'trend' in chart_paths:
        add_image(doc, chart_paths['trend'], f'图1 {county}路面状况指数年度趋势')
    
    # 添加年度趋势表格
    if table_data and 'yearly_trend' in table_data:
        trend_df = table_data['yearly_trend']
        if not trend_df.empty:
            add_heading(doc, '2.1 历年指数均值汇总', 2)
            add_summary_table(doc, trend_df, f'表1 {county}历年路面状况指数均值')
    
    # 生成AI趋势分析文字
    trend_text = None
    if llm and table_data and 'yearly_trend' in table_data:
        try:
            from src.llm_writer import generate_trend_analysis_text
            trend_data = table_data['yearly_trend'].to_dict() if hasattr(table_data['yearly_trend'], 'to_dict') else table_data['yearly_trend']
            trend_text = generate_trend_analysis_text(trend_data, county, llm)
        except Exception as e:
            print(f"生成趋势分析失败: {e}")
    
    if trend_text:
        add_body_text(doc, trend_text)
    else:
        add_body_text(doc, '根据历年检测数据，路面技术状况指数呈现稳定趋势。')
    
    # ═══════════════════════
    # 三、路面技术状况等级分布
    # ═══════════════════════
    _progress("正在生成等级分布分析...")
    add_heading(doc, '三、路面技术状况等级分布分析', 1)
    
    # 添加等级分布图表
    if chart_paths and 'grade_dist' in chart_paths:
        add_image(doc, chart_paths['grade_dist'], f'图2 {county} {latest_year}年PQI等级里程分布')
    
    if chart_paths and 'grade_stacked' in chart_paths:
        add_image(doc, chart_paths['grade_stacked'], f'图3 {county}历年PQI等级里程比例分布')
    
    # 添加等级分布表格
    if table_data and 'grade_dist' in table_data:
        grade_df = table_data['grade_dist']
        if not grade_df.empty:
            add_heading(doc, '3.1 PQI等级里程分布', 2)
            add_summary_table(doc, grade_df, f'表2 {county} {latest_year}年PQI等级里程分布')
    
    # 生成AI等级分析文字
    grade_text = None
    if llm and table_data and 'grade_dist' in table_data:
        try:
            from src.llm_writer import generate_grade_analysis_text
            grade_data = table_data['grade_dist'].to_dict() if hasattr(table_data['grade_dist'], 'to_dict') else table_data['grade_dist']
            grade_text = generate_grade_analysis_text(grade_data, county, latest_year, llm)
        except Exception as e:
            print(f"生成等级分析失败: {e}")
    
    if grade_text:
        add_body_text(doc, grade_text)
    else:
        add_body_text(doc, '路面技术状况等级分布基本合理，大部分路段处于良好状态。')
    
    # 路面类型和技术等级
    add_heading(doc, '3.2 路面类型与技术等级分布', 2)
    
    if chart_paths and 'pavement_type' in chart_paths:
        add_image(doc, chart_paths['pavement_type'], f'图4 {county} {latest_year}年路面类型里程分布', width_inches=4.5)
    
    if chart_paths and 'tech_grade' in chart_paths:
        add_image(doc, chart_paths['tech_grade'], f'图5 {county} {latest_year}年技术等级里程分布', width_inches=4.5)
    
    add_body_text(doc, '路面类型以沥青路面为主，技术等级以二级和三级公路为主。')
    
    # ═══════════════════════
    # 四、养护需求分析
    # ═══════════════════════
    _progress("正在生成养护需求分析...")
    add_heading(doc, '四、养护需求分析与资金估算', 1)
    
    # 添加养护需求图表
    if chart_paths and 'maintenance' in chart_paths:
        add_image(doc, chart_paths['maintenance'], f'图6 {county} {latest_year}年养护需求里程与费用估算')
    
    # 添加养护需求表格
    if table_data and 'maintenance' in table_data:
        maint_df = table_data['maintenance']
        if not maint_df.empty:
            add_heading(doc, '4.1 养护需求汇总', 2)
            add_summary_table(doc, maint_df, f'表3 {county} {latest_year}年养护需求汇总（费用参考单价：路面改造约300元/m²，预防养护约150元/m²）')
    
    # 生成AI养护分析文字
    maint_text = None
    if llm and table_data and 'maintenance' in table_data:
        try:
            from src.llm_writer import generate_maintenance_text
            total_cost = table_data.get('maintenance_cost', 0)
            maint_text = generate_maintenance_text(table_data['maintenance'], county, latest_year, total_cost, llm)
        except Exception as e:
            print(f"生成养护分析失败: {e}")
    
    if maint_text:
        add_body_text(doc, maint_text)
    else:
        add_body_text(doc, '根据路面技术状况检测结果，部分路段需要进行养护处理。')
    
    # ═══════════════════════
    # 五、养护需求预测
    # ═══════════════════════
    _progress("正在生成养护需求预测...")
    add_heading(doc, f'五、养护需求预测（{latest_year+1}—{latest_year+5}年）', 1)
    
    # 添加预测图表
    if chart_paths and 'prediction' in chart_paths:
        add_image(doc, chart_paths['prediction'], f'图7 {county}未来5年养护需求预测')
    
    # 添加预测表格
    if table_data and 'prediction' in table_data:
        pred_df = table_data['prediction']
        if not pred_df.empty:
            add_heading(doc, '5.1 年度养护需求预测', 2)
            add_summary_table(doc, pred_df, f'表4 {county}未来5年养护需求预测')
    
    add_heading(doc, '5.2 预测原则', 2)
    for item in [
        '数据驱动原则：以最新年度路面技术状况检测数据为基础，采用定量分析方法进行预测。',
        '分路施策原则：根据不同路线、不同路面类型、不同技术等级，分别建立预测模型。',
        '动态滚动原则：预测周期为5年，可根据每年检测结果动态更新。',
    ]:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    add_heading(doc, '5.3 预测模型说明', 2)
    model_text = '本报告采用指数衰减模型预测PQI、PCI、RQI未来5年的变化，其中年衰减率依据《普通国省干线路面性能衰变研究》等现有研究成果标定。养护触发条件参照路面技术状况指标预测值，判断路段在当年是否需要养护及养护类型：路面改造（PQI/PCI<60）、中修（PQI/PCI 60-70）、预防养护（PQI/PCI 70-80）、日常养护（PQI/PCI≥80）。'
    add_body_text(doc, model_text)
    
    # 生成AI预测分析文字
    forecast_text = None
    if llm and table_data and 'yearly_trend' in table_data:
        try:
            from src.llm_writer import generate_forecast_text
            trend_data = table_data['yearly_trend'].to_dict() if hasattr(table_data['yearly_trend'], 'to_dict') else table_data['yearly_trend']
            forecast_text = generate_forecast_text(county, latest_year, trend_data, llm)
        except Exception as e:
            print(f"生成预测分析失败: {e}")
    
    if forecast_text:
        add_body_text(doc, forecast_text)
    
    # ═══════════════════════
    # 结论与建议
    # ═══════════════════════
    _progress("正在生成结论与建议...")
    add_heading(doc, '六、结论与建议', 1)
    
    # 生成AI结论文字
    conclusion_text = None
    if llm:
        try:
            from src.llm_writer import generate_conclusion_text
            conclusion_text = generate_conclusion_text(county, stats, llm)
        except Exception as e:
            print(f"生成结论失败: {e}")
    
    if conclusion_text:
        add_body_text(doc, conclusion_text)
    else:
        add_body_text(doc, '通过对路面技术状况的分析，建议加强对重点路段的养护管理，合理安排养护资金，确保公路安全畅通。')
    
    # ═══════════════════════
    # 附件说明
    # ═══════════════════════
    doc.add_page_break()
    add_heading(doc, '附件：数据说明', 1)
    notes = [
        'PQI（路面质量指数）：综合评价路面整体技术状况，满分100分，≥90为优，80-90为良，70-80为中，60-70为次，<60为差。',
        'PCI（路面状况指数）：评价路面破损状况，分级标准与PQI相同。',
        'RQI（行驶质量指数）：评价路面行驶舒适性，≥95为优，90-95为良，85-90为中，80-85为次，<80为差。',
        '本报告数据来源于历年公路路况普查检测数据，已剔除上报剔除路段和重复路段。',
        f'养护费用估算为参考值，实际费用需根据现场勘查和设计确定。',
    ]
    for note in notes:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'● {note}')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = '仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 保存
    _progress("正在保存报告...")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    preview_text += f"\n\n报告已保存到：{output_path}"
    return preview_text
